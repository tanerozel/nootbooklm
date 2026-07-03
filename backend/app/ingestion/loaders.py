"""Document loaders for PDF, DOCX, plain text/markdown, and web URLs."""
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

import requests
from bs4 import BeautifulSoup


class LoadedDocument:
    """Holds extracted text with per-page segments and metadata."""

    def __init__(self, pages: list[dict], title: str, source_type: str):
        # pages: list of {"page_number": int, "text": str}
        self.pages = pages
        self.title = title
        self.source_type = source_type


def load_pdf(file_path: str, title: Optional[str] = None) -> LoadedDocument:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages.append({"page_number": i + 1, "text": text.strip()})

    return LoadedDocument(
        pages=pages,
        title=title or Path(file_path).stem,
        source_type="pdf",
    )


def load_docx(file_path: str, title: Optional[str] = None) -> LoadedDocument:
    from docx import Document

    doc = Document(file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return LoadedDocument(
        pages=[{"page_number": None, "text": full_text}],
        title=title or Path(file_path).stem,
        source_type="docx",
    )


def load_text(file_path: str, title: Optional[str] = None) -> LoadedDocument:
    text = Path(file_path).read_text(encoding="utf-8", errors="replace")
    return LoadedDocument(
        pages=[{"page_number": None, "text": text}],
        title=title or Path(file_path).stem,
        source_type="text",
    )


def load_url(url: str, title: Optional[str] = None) -> LoadedDocument:
    headers = {"User-Agent": "Mozilla/5.0 (NootbookLM ingestion bot)"}
    resp = requests.get(url, headers=headers, timeout=30)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    # Remove scripts, styles, nav
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    page_title = title or (soup.title.string.strip() if soup.title else url)

    return LoadedDocument(
        pages=[{"page_number": None, "text": text}],
        title=page_title,
        source_type="url",
    )


def load_document(
    source_type: str,
    file_path: Optional[str] = None,
    url: Optional[str] = None,
    title: Optional[str] = None,
) -> LoadedDocument:
    if source_type == "pdf":
        return load_pdf(file_path, title=title)
    if source_type == "docx":
        return load_docx(file_path, title=title)
    if source_type in ("text", "txt", "md", "markdown"):
        return load_text(file_path, title=title)
    if source_type == "url":
        return load_url(url, title=title)
    raise ValueError(f"Unsupported source type: {source_type}")
